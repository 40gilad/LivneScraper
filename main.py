import requests
from bs4 import BeautifulSoup as BS
from googlesearch import search
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from dotenv import load_dotenv
import os
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
import datetime

# region Globals

GPT3 = "gpt-3.5-turbo"
GPT4 = "gpt-4-0125-preview"
document = Document()
section = document.sections[0]
section.right_to_left = True
output_folder = None
bizportal_root = "https://www.bizportal.co.il"
bizportal_data_link = "https://www.bizportal.co.il/bonds/quote/bondsdata/"
bizportal_search_page = "https://www.bizportal.co.il/list/searchpapers?search="
chrome_driver_path = r".\chromedriver.exe"
driver = None
# endregion

# region Load .env


env_path = r'C:\Users\40gil\Desktop\Helpful\Scraping\LivneScraperPY\LivneScraper.env'
try:
    load_dotenv(dotenv_path=env_path)
    api_key = os.getenv('OPENAI_KEY')
    if api_key is None:
        raise ValueError("OpenAI API key not found in environment variables.")
except Exception as err:
    print(f"Error: {err}")
    # Handle the error accordingly, e.g., exit the script or provide a default API key.


# endregion

# region Doc manipulation

def add_paragraph(text=None, style=None, subheadline=False):
    if text is None:
        return
    par = document.add_paragraph(text, style=style)
    run = par.runs[0]  # Assuming there is only one run in the paragraph
    font = run.font
    font.name = 'Calibri'  # Set the font name
    if subheadline:
        font.size = Pt(18)  # Set the font size
    else:
        font.size = Pt(12)  # Set the font size

    par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def add_headline(text, size=2):
    global document
    hed = document.add_heading(text, size)
    hed.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def insert_bizportal_data(root_bonds_url_data=None):
    global driver
    if root_bonds_url_data is None or driver is None:
        raise ValueError("root_bonds_url_data or driver are None")
        return
    sub_headline = f'{len(root_bonds_url_data)}: אגח '
    add_paragraph(text=f"{sub_headline}\n", subheadline=True)

    for d in root_bonds_url_data:

        # get root bond page
        driver.get(d['root_link'])
        page_source = driver.page_source
        soup = BS(page_source, 'html.parser')

        # add bonds sum as subheadline and the name of the bond
        add_paragraph(text=f" - {d['name']}  ", style='List Bullet', subheadline=True)

        # from root page
        spans = soup.find_all(name='span', class_='label')

        # every if inc counter, so if all ifs happens -> break
        iter_counter = 0

        for span in spans:

            if iter_counter == 2:  # already got into all ifs
                iter_counter = 0
                break
            curr_span = span.get_text()

            # todo NEED TO INSERT VALUE

            if curr_span == 'ענף':

                iter_counter += 1

                add_paragraph(text=f"{span.find_next('span').get_text()} ", style='List Bullet 2')
            elif curr_span == 'מח"מ:':

                iter_counter += 1

                val = span.find_next('span').get_text()
                text_to_insert = f"{curr_span} {val} "
                add_paragraph(text=text_to_insert, style='List Bullet 2')

        # get data bond page
        driver.get(d['data_link'])
        page_source = driver.page_source
        soup = BS(page_source, 'html.parser')

        # from data page
        spans = soup.find_all(name='td', class_='label')
        for span in spans:

            if iter_counter == 3:
                # unknwon ribit type and ribit val order of reading, so insrting at the end
                add_paragraph(text=f"{ribit_type} {ribit_val}", style='List Bullet 2')
                break

            curr_span = span.get_text()
            if curr_span == 'סוג ריבית':
                ribit_type = f" ריבית {span.find_next(name='td', class_='num').get_text()} "
                iter_counter += 1

            elif curr_span == 'שיעור ריבית':

                iter_counter += 1
                val = span.find_next(name='td', class_="num").get_text()
                ribit_val = f"% {val} "

            elif curr_span == 'דרוג מידרוג + אופק':

                iter_counter += 1
                text_to_insert = "דירוג מידרוג "
                text_to_insert += span.find_next(name='td', class_="num").get_text()
                add_paragraph(text=text_to_insert, style='List Bullet 2')


# endregion

# region Scraping helpers
def format_html_txt(html_string):
    soup = BS(html_string, 'html.parser')
    text_without_tags = soup.get_text()
    return text_without_tags


def format_wiki_text(wiki_txt):
    txt = ""
    for p in wiki_txt:
        txt += f'{p.get_text()}\n'
    while '[' in txt:
        indx = txt.index('[')
        txt = txt[:indx] + txt[indx + 3:]  # " bla bla[1] bla bla"
    return txt


def get_link(company, wiki=False, maya=False, bizportal=False,
             globs=False, bizpotal_juice=False, themarker=False, calcalist=False,
             facebook=False, instagram=False, linkedin=False):
    # TODO: keep on with juice sites
    if wiki:
        query = f"{company} ויקיפדיה "
    elif maya:
        query = f"{company} מאיה "
    elif bizportal:
        query = f"{company} ביזפורטל אגח "
    elif globs:
        query = f"{company} גלובס "
    elif bizpotal_juice:
        query = f"{company} ביזפורטל כל הכתבות והמידע הקשורים "
    elif calcalist:
        query = f"{company} כלכליסט "
    elif instagram:
        query = f"{company} instagram"
    elif facebook:
        query = f"{company} facebook"
    elif linkedin:
        query = f"{company} linkedin"

    search_res = search(query, tld="co.il", stop=3)
    for j in search_res:
        if (
                (wiki and 'he.wikipedia.org/wiki/' in j)
                or (maya and 'maya.tase.co.il' in j)
                or (bizportal and 'bizportal.co.il' in j)
                or (bizpotal_juice and 'www.bizportal.co.il/list/tags/' in j)
                or (globs and 'globes.co.il/news/' in j and j.endswith('.tag'))
                or (themarker and 'themarker.com/ty-tag/' in j)
                or (calcalist and 'www.calcalist.co.il/tags' in j)
                or (facebook and 'facebook.com' in j)
                or (instagram and 'instagram.com' in j)
                or (linkedin and 'linkedin.com' in j)

        ):
            return j.split('?')[0]
    return None


def get_bonds(bond_name):
    """

    :param bond_name:
    :param driver:
    :return: [{
                'root_link': root_link,
                'name': bond_name,
                'bond_number': bond_number
            },
            {
                .....
            }
            ]
    """
    global driver

    bond_name = bond_name.replace(" ", "%20")
    link = f"{bizportal_search_page}{bond_name}"
    driver.get(link)
    page_source = driver.page_source
    soup = BS(page_source, 'html.parser')
    bonds_links = soup.find_all('a', class_='link')
    ret = []
    for link in bonds_links:
        root_link = f"{bizportal_root}{link.get('href')}"
        root_splitted = root_link.split("/")
        bond_number = root_splitted[len(root_splitted) - 1]
        ret.append(
            {
                'root_link': root_link,
                'name': link.get_text(),
                'bond_number': bond_number,
                'data_link': f"{bizportal_data_link}{bond_number}"
            }
        )
    return ret


def crawl_reports(reports=None):
    if reports:
        for i in range(0, 10):
            txt = '- ' + reports[i].find_all('span', class_='feedItemDateMobile ng-binding')[0].get_text(separator=' ',
                                                                                                         strip=True)
            if txt is not None:
                add_paragraph(txt, style='List Bullet')
            txt = reports[i].find_all('a', class_='messageContent')[0].get_text(separator=' ', strip=True) + '\n'

            if txt is not None:
                add_paragraph(f"    {txt}")


def get_data_from_site(link=None, bond_name=None, wiki=False, wiki_paragraphs=2,
                       maya=False, key_people=False, bizportal=False, maya_reports=False,
                       globs=False, bizportal_juice=False, themarker=False, calcalist=False, juice_articles_amount=3):
    global driver

    if not bizportal:
        # the websites that demand driver.
        driver.get(link)
        page_source = driver.page_source
        soup = BS(page_source, 'html.parser')

    if wiki:
        wiki_data = []
        for i in range(wiki_paragraphs - 1):
            wiki_data.append(soup.find('p'))
            wiki_data.append(wiki_data[i].find_next('p'))
        return f'\n {format_wiki_text(wiki_data)}'

    elif maya:
        share_holders_table = soup.find_all(name="div", class_="listTable share-holders-grid")
        rows = share_holders_table[0].find_all(name="div", class_="ng-scope tableRow")
        ret_txt = ''
        for row in rows:
            try:
                comp_name_tag = row.find(name="div")
                precentage_tag = row.find(name="div", class_="tableCol col_6 ng-binding ng-scope")
                if comp_name_tag is not None and precentage_tag is not None:
                    comp_name = comp_name_tag.get_text()
                    precentage = precentage_tag.get_text()
                    ret_txt += f"% {comp_name} - {precentage}\n"
            except AttributeError:
                return ret_txt

        return ret_txt

    elif key_people:
        key_people_table = soup.find_all(name="div", class_="mobileTableFrame")[2]
        rows = key_people_table.find_all(name="div", class_="ng-scope tableRow")

        ret_txt = ''
        for row in rows:
            person_name_tag = row.find(name="div")
            person_role = person_name_tag.find_next(name="div").get_text()
            if person_role != "דירקטור רגיל":
                person_name = person_name_tag.get_text()
                if "מאזן" in person_name:
                    return "לא נמצאו אנשי מפתח"
                ret_txt += f"{person_role} - {person_name}\n"
        return ret_txt

    elif maya_reports:
        try:
            reports_section = soup.find('maya-reports')
            reports = reports_section.find_all('div', class_='feedItem ng-scope')
            crawl_reports(reports=reports)

        except Exception as e:
            print(f"{link}, for maya reports\nexception: {e}")
            return 'משהו השתבש'

    elif bizportal:
        root_bonds_url_data = get_bonds(bond_name=bond_name)
        insert_bizportal_data(root_bonds_url_data=root_bonds_url_data)

    elif globs:
        globs_root="https://www.globes.co.il/"
        arts = soup.find_all('div', class_='tagit')
        for i in range(0, juice_articles_amount):
            art = arts[i]
            title = art.find('h3', class_='tagit__title').get_text()
            href = art.find('a', class_='tagit__link').get('href')
            description = art.find('p', class_='tagit__subtitle').get_text()
            add_paragraph(title,subheadline=True)
            add_paragraph(description)
            add_paragraph(f'{globs_root}/{href}\n')
        return None

    elif bizportal_juice:
        art = soup.find('div', class_='MainArticleBox')#first big article
        title=art.find(name='div', class_='text').get_text()
        href = art.find('a', class_='img-link').get('href')
        add_paragraph(f'{title} - {href}')
        arts=soup.find_all(name='ul',id="secondary-article-two")[0] # two side articles
        for art in arts.find_all('li', class_='bordered-item-two'):
            title_and_href=art.find_all('a')[3]
            href=title_and_href.get('href')
            title=title_and_href.get_text()
            add_paragraph(f'{title} - {href}')
        return None


    elif themarker:
        arts=soup.find_all('article', _class='dy x d bm gd ge gf gg gh gi gj gk gl gm gn go')




    elif calcalist:
        pass

    return None


# endregion

# region GEPETO

def ask_gepeto(prompt, model=GPT3):
    # ChatGPT API endpoint
    endpoint = "https://api.openai.com/v1/chat/completions"

    # Headers containing the Authorization with your API key
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }

    # Data payload containing the prompt and other parameters
    data = {
        "model": model,  # You can choose different models as per your requirement
        "messages": [
            {"role": "user", "content": prompt}
        ]
    }

    try:
        # Making a POST request to the API
        response = requests.post(endpoint, json=data, headers=headers)

        # Check if the request was successful
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        else:
            return f"Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"Error: {str(e)}"


# endregion

# region add_<section>_sum
def add_wiki_sum(wiki_link=None):
    if wiki_link is not None:
        wiki_text = get_data_from_site(link=wiki_link, wiki=True)
    else:
        wiki_text = 'לא נמצא לינק לויקיפדיה'
    add_paragraph(wiki_text)


def add_maya_sum(maya_link=None):
    maya_text = ''
    if maya_link is not None:
        try:
            maya_text = get_data_from_site(link=maya_link, maya=True)
        except Exception as e:
            raise Exception(f"add_maya_sum Error: {str(e)}")
            maya_text = 'קרתה תקלה. נסה שוב '
    else:
        maya_text = 'לא נמצא לינק למאיה'
    add_paragraph(maya_text)


def add_bizportal_sum(company_name=None, bond_name=None):
    bizportal_text = ''
    if company_name is not None:
        try:
            get_data_from_site(bond_name=bond_name, bizportal=True)
        except Exception as e:
            raise Exception(f"add_bizportal_sum Error: {str(e)}")


def add_key_people(maya_link=None):
    key_text = ''
    if maya_link is not None:
        try:
            key_text = get_data_from_site(link=maya_link, key_people=True)
        except Exception as e:
            raise Exception(f"add_bizportal_sum Error: {str(e)}")
            key_text = 'קרתה תקלה. נסה שוב '
    else:
        key_text = 'לא נמצא לינק למאיה'
    add_paragraph(key_text)


def add_last_reports(maya_link=None):
    maya_text = ''
    if maya_link is not None:
        try:
            maya_text = get_data_from_site(link=maya_link, maya_reports=True)
        except Exception as e:
            raise Exception(f"add_bizportal_sum Error: {str(e)}")
            key_text = 'קרתה תקלה. נסה שוב '
    else:
        maya_text = 'לא נמצאו דיווחים אחרונים במאיה'
        add_paragraph(maya_text)


def add_juice(company_name=None,news_websites=None):
    if company_name is None:
        return "שגיאה בשם החברה"
    for key in news_websites:

        link = None
        juice_text = None

        # if key == 'globs':
        #     link = get_link(company=company_name, globs=True)
        #     if link is not None:
        #         add_paragraph(text=f'{key}:\n', style="List Bullet",subheadline=True)
        #         juice_text = get_data_from_site(link=link, globs=True)

        # elif key == 'bizportal':
        #     link = get_link(company=company_name, bizpotal_juice=True)
        #     if link is not None:
        #         add_paragraph(text=f'{key}:\n', style="List Bullet",subheadline=True)
        #         juice_text = get_data_from_site(link=link, bizportal_juice=True)

        if key == 'themarker':
            link = f'https://www.themarker.com/search-results?q={company_name.replace(" ", "+")}'
            if link is not None:
                add_paragraph(text=f'{key}:\n', style="List Bullet",subheadline=True)
                juice_text = get_data_from_site(link=link, themarker=True)
        #
        # elif key == 'calcalist':
        #     link = get_link(company=company_name, calcalist=True)
        #     if link is not None:
        #         add_paragraph(text=f'{key}:\n', style="List Bullet",subheadline=True)
        #         juice_text = get_data_from_site(link=link, calcalist=True)

        if juice_text is not None:
            add_paragraph(juice_text)


def add_social(company_name=None):
    # ----------- FACEBOOK URL -----------#
    txt = '-פייסבוק '
    add_paragraph(txt, style='List Bullet')
    link = get_link(company_name, facebook=True)
    if link is not None:
        add_paragraph(f"    {link}")

    # ----------- INSTAGRAM URL -----------#
    txt = '-אינסטגרם '
    add_paragraph(txt, style='List Bullet')
    link = get_link(company_name, instagram=True)
    if link is not None:
        add_paragraph(f"    {link}")

    # ----------- LINKEDIN URL -----------#
    txt = '-לינקדאין '
    add_paragraph(txt, style='List Bullet')
    link = get_link(company_name, linkedin=True)
    if link is not None:
        add_paragraph(f"    {link}")


# endregion


def scrape_and_sum(_company_name=None, bond_name=None):
    if _company_name is None or bond_name is None:
        raise ValueError('Scrape_and_sum: _Company_name or bond_name are None. both should have valid value')

    global driver

    wiki_headline = '    כללי: ויקיפדיה'
    maya_headline = 'בעלות: מאיה'
    bizportal_headline = 'ני"ע: ביזפורטל'
    key_people_headline = 'אנשי מפתח : מאיה '
    imeidiate_reports = 'דיווחים מידיים: מאיה'
    juice_headline = 'עיתונות: '
    social_headline = ':סושיאל'
    add_headline(text=_company_name, size=0)

    # ----------- WIKIPEDIA DATA -----------#
    # add_headline(text=wiki_headline)
    # wiki_link = get_link(company=_company_name, wiki=True)
    # add_wiki_sum(wiki_link=wiki_link)
    # # -------------- MAYA DATA --------------#
    # add_headline(text=maya_headline)
    # maya_link = get_link(company=_company_name, maya=True)
    # add_maya_sum(maya_link=maya_link)
    # # -------------- BIZPORTAL DATA --------------#
    # add_headline(text=bizportal_headline)
    # bizportal_link = get_link(company=_company_name, bizportal=True)
    # add_bizportal_sum(company_name=_company_name, bond_name=bond_name)
    # # -------------- KEY PEOPLE DATA --------------#
    # add_headline(text=key_people_headline)
    # add_key_people(maya_link=maya_link)
    # # -------------- IMMIDIATE REPORTS --------------#
    # add_headline(text=imeidiate_reports)
    # link = f"{maya_link}?view=reports&q=%7B%22DateFrom%22:%222023-02-13T22:00:00.000Z%22,%22DateTo%22:%222024-02-13T22:00:00.000Z%22,%22Page%22:1,%22entity%22:%221840%22,%22events%22:%5B%5D,%22subevents%22:%5B%5D%7D"
    # add_last_reports(maya_link=link)
    # -------------- JUICE DATA --------------#
    add_headline(text=juice_headline)
    add_juice(company_name=_company_name,news_websites=["globs","themarker","calcalist","bizportal"])
    # -------------- SOCIAL DATA --------------#
    add_headline(text=social_headline)
    add_social(company_name=_company_name)

    driver.quit()


def start(_company_name=None, bond_name=None, to_save_path=None, chrome_driver=None):
    global output_folder, chrome_driver_path, driver
    if _company_name is None or bond_name is None or to_save_path is None or chrome_driver is None:
        raise ValueError('in start function, you must provide all necessary arguments')
    output_folder = to_save_path
    chrome_driver_path = chrome_driver
    driver = webdriver.Chrome(service=Service(executable_path=chrome_driver_path))

    try:
        scrape_and_sum(_company_name=_company_name, bond_name=bond_name)
    except Exception as e:
        raise Exception(f'in start function: {e}')

    if output_folder is None:
        try:
            document.save(f'{_company_name}.docx')
        except:
            document.save(f'{_company_name}_{datetime.datetime.now().strftime("%d.%m_%H.%M")}.docx')
    else:
        try:
            document.save(f'{output_folder}\\{_company_name}.docx')
        except:
            document.save(f'{output_folder}\\{_company_name}_{datetime.datetime.now().strftime("%d.%m_%H.%M")}.docx')


if __name__ == '__main__':
    paragraph = document.add_paragraph()
    company_name = 'בנק הפועלים'
    driver = webdriver.Chrome(service=Service(executable_path=chrome_driver_path))
    scrape_and_sum(_company_name=company_name, bond_name='דליה אגח')
    # try:
    #     start()
    # except Exception as e:
    #     print(f"\nAn error occurred: " + str(e))
    #     exit(1)
